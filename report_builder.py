"""Build DOCX reports from CSWL section CSVs."""

from __future__ import annotations

import os

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from metrics_core import compute_subsection_bandwise_metrics

CSWL_SECTION_ORDER = [
    "Z Scored FFT Absolute Power",
    "Z Scored FFT Power Ratio",
    "Z Scored Peak Frequency",
]

BULLET_METRICS = [
    "Absolute Sum:",
    "Average Absolute Value:",
    "Delta:",
    "Percent Change:",
    "Total Rows:",
    '"Normalize = Yes":',
    '"Normalize = No":',
    '"Normalize = NS":',
]


def create_combined_document(csv_files: list[str], output_path: str) -> None:
    """Create a single DOCX combining all section CSV files with band-wise metrics."""
    doc = Document()
    eeg_heading = doc.add_paragraph()
    run = eeg_heading.add_run("EEG//")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_heading("Band-wise Metrics Analysis: ", level=1)

    section_to_file = {
        os.path.basename(f).replace(".csv", ""): f for f in csv_files
    }

    for section in CSWL_SECTION_ORDER:
        file_path = section_to_file.get(section)
        if not file_path:
            continue

        df = pd.read_csv(file_path)
        metrics_text = compute_subsection_bandwise_metrics(
            df,
            subsection_col="Subsection",
            band_col="Band",
            set1_col="T1 Z",
            set2_col="T2 Z",
            normalize_col="Normalize",
        )
        doc.add_heading(section, level=2)
        for block in metrics_text.strip().split("\n Subsection: "):
            if not block.strip():
                continue
            if not block.startswith("Subsection:"):
                block = "Subsection: " + block
            lines = block.strip().split("\n", 1)
            subsection_title = lines[0].replace("Subsection: ", "").strip()
            content = lines[1] if len(lines) > 1 else ""
            para = doc.add_paragraph()
            run = para.add_run(f"Subsection: {subsection_title}")
            run.bold = True
            for line in content.strip().split("\n"):
                line_strip = line.strip()
                if any(
                    line_strip.lstrip().startswith(metric) for metric in BULLET_METRICS
                ):
                    p = doc.add_paragraph(f"• {line_strip}", style="Normal")
                    p.paragraph_format.left_indent = Inches(0.4)
                elif line_strip:
                    doc.add_paragraph(line_strip, style="Normal")
        doc.add_paragraph()

    doc.save(output_path)
