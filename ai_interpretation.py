"""
AI interpretation of QEEG metrics using Gemini API.
Supports vertical (Set 1 vs Set 2) and horizontal (timeline / matrix) analysis modes.
"""

from __future__ import annotations

import logging
import os
import re
from typing import Any, Dict, List, Optional, Tuple

import requests
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

from qeeg_statistical_analysis import AnalysisMode, AnalysisPackage, run_statistical_analysis

load_dotenv()

logger = logging.getLogger(__name__)

MAX_INPUT_CHARS = 28000
GEMINI_ENDPOINT = (
    "https://generativelanguage.googleapis.com/v1beta/models/"
    "{model}:generateContent?key={api_key}"
)


def _strip_limitations_section(text: str) -> str:
    if not text:
        return text
    return re.sub(
        r"(?ms)^##\s*Limitations\s+and\s+caveats\s*\n.*",
        "",
        text,
    ).strip()


def _build_prompt(pkg: AnalysisPackage) -> str:
    content = pkg.to_prompt_context()[:MAX_INPUT_CHARS]
    if len(pkg.to_prompt_context()) > MAX_INPUT_CHARS:
        content += "\n\n[Note: Input was truncated for model context limits.]"

    if pkg.mode == "vertical":
        intro = (
            "You are assisting with **QEEG Vertical Analysis**: direct EC (Eyes Closed) vs "
            "EO (Eyes Open) comparison from CSWL extraction.\n\n"
            "Focus on:\n"
            "- Band-wise and subsection aggregates (LEFT / RIGHT / CENTER)\n"
            "- Percent change, normalize direction counts (Yes/No/NS)\n"
            "- Abnormal z-score rates (|z| > 1.96) per condition\n"
            "- Largest regional shifts between EC and EO\n\n"
        )
    else:
        intro = (
            "You are assisting with **QEEG Horizontal Analysis**: timeline trajectory, "
            "matrix statistics, and symptom-pattern scores across EC → EO.\n\n"
            "Focus on:\n"
            "- Mean, SD, median, IQR, and % abnormal across EC and EO per segment/band\n"
            "- Trend direction from EC → EO\n"
            "- Symptom-engine pattern scores (non-diagnostic correlates)\n"
            "- Regional contributors when listed\n\n"
        )

    return (
        intro
        + "This is **not** a diagnosis. Use cautious clinical language and cite numbers from the data.\n\n"
        "Use EXACT headings:\n"
        "## Executive Summary\n"
        "## Descriptive overview by EEG section\n"
        "## Set 1 vs Set 2 — comparative highlights\n"
        "## Cross-section patterns\n"
        "## Clinical considerations (non-diagnostic)\n\n"
        "Rules:\n"
        "- Do not invent values absent from the data.\n"
        "- Prefer quantitative cues from COMPUTED DIAGNOSTICS and METRICS DIGEST.\n\n"
        "--- DATA ---\n\n"
        f"{content}"
    )


def _extract_text(payload: Dict[str, Any]) -> str:
    candidates = payload.get("candidates") or []
    for c in candidates:
        parts = ((c.get("content") or {}).get("parts") or [])
        text = "\n".join([(p.get("text") or "").strip() for p in parts]).strip()
        if text:
            return text
    return ""


def invoke_gemini_analysis(prompt: str) -> str:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY is not set in .env")

    model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
    url = GEMINI_ENDPOINT.format(model=model, api_key=api_key)
    body = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.7, "topP": 0.95},
    }
    resp = requests.post(url, json=body, timeout=120)
    if not resp.ok:
        raise ValueError(f"Gemini API error: {resp.status_code} {resp.text[:400]}")
    text = _extract_text(resp.json())
    if not text:
        raise ValueError("Gemini API returned empty analysis text")
    return _strip_limitations_section(text)


def generate_analysis_report(
    csv_files: List[str],
    mode: AnalysisMode,
    *,
    set1_label: str = "Set 1",
    set2_label: str = "Set 2",
) -> Tuple[AnalysisPackage, str]:
    """Run statistics + Gemini; returns (package, markdown report)."""
    pkg = run_statistical_analysis(
        csv_files,
        mode,
        set1_label=set1_label,
        set2_label=set2_label,
    )
    prompt = _build_prompt(pkg)
    report = invoke_gemini_analysis(prompt)
    return pkg, report


def write_analysis_docx(path: str, title: str, report_markdown: str) -> None:
    """Write Gemini narrative to DOCX."""
    doc = Document()
    heading = doc.add_paragraph()
    run = heading.add_run("EEG//")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_heading(title, level=1)

    for line in report_markdown.split("\n"):
        line = line.rstrip()
        if not line.strip():
            continue
        h2 = re.match(r"^##\s+(.+)$", line)
        h3 = re.match(r"^###\s+(.+)$", line)
        if h2:
            doc.add_heading(h2.group(1).strip(), level=2)
        elif h3:
            doc.add_heading(h3.group(1).strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(path)


def try_generate_analysis_docx(
    output_path: str,
    csv_files: List[str],
    mode: AnalysisMode,
    *,
    set1_label: str = "Set 1",
    set2_label: str = "Set 2",
) -> Tuple[bool, Optional[str], Optional[str]]:
    """Returns (success, error_message, report_markdown)."""
    if not os.getenv("GEMINI_API_KEY"):
        return False, "GEMINI_API_KEY not set in .env", None
    try:
        _, report = generate_analysis_report(
            csv_files, mode, set1_label=set1_label, set2_label=set2_label
        )
        title = (
            "QEEG Vertical (EC vs EO) — AI Report"
            if mode == "vertical"
            else "QEEG Horizontal (timeline / matrix) — AI Report"
        )
        write_analysis_docx(output_path, title, report)
        return True, None, report
    except Exception as e:
        logger.exception("AI analysis failed")
        return False, str(e), None
