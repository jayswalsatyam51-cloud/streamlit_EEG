from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import os
import shutil
from typing import List
import tempfile
from datetime import datetime
import json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob
import numpy as np
import zipfile
from dotenv import load_dotenv
import uuid
import html

load_dotenv()

# Import the extraction functions from pdf_extraction.py
from pdf_extraction import extract_coherence_phase_lag, extract_other_sections, merge_data, save_to_csv
from upload_utils import upload_to_s3
from metrics_core import compute_subsection_bandwise_metrics
from ai_interpretation import try_generate_ai_interpretation_docx
from cswl_pipeline import (
    MAX_CSWL_UPLOAD_BYTES,
    process_cswl_pair_to_output_dir,
    secure_cswl_filename,
)

app = FastAPI(
    title="PDF Data Extraction API",
    description="API for extracting and processing data from PDF files",
    version="1.0.0"
)

# Add CORS middleware (aligned with VNG-lite FastAPI configuration)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create a directory to store uploaded files
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

STATIC_DIR = Path(__file__).resolve().parent / "static"
if STATIC_DIR.is_dir():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


@app.get("/ui", response_class=HTMLResponse, include_in_schema=False)
async def web_ui():
    """Browser UI for PDF upload and result downloads (same-origin API by default)."""
    index = STATIC_DIR / "index.html"
    if index.is_file():
        return FileResponse(str(index))
    raise HTTPException(status_code=404, detail="UI not found")

@app.post("/extraction")
async def extract_data(
    pdf1: UploadFile = File(...),
    pdf2: UploadFile = File(...)
):
    """
    Upload 2 PDFs, extract data, perform all calculations, and return the processed results.
    """
    # Validate file types
    for pdf in [pdf1, pdf2]:
        if not pdf.filename.lower().endswith('.pdf'):
            raise HTTPException(
                status_code=400, 
                detail=f"File {pdf.filename} is not a PDF"
            )
    
    try:
        # Create a temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_paths = []
            
            # Save uploaded files
            for pdf in [pdf1, pdf2]:
                file_path = os.path.join(temp_dir, pdf.filename)
                with open(file_path, "wb") as buffer:
                    shutil.copyfileobj(pdf.file, buffer)
                pdf_paths.append(file_path)
            
            # Process the PDFs
            coherence_phase_lag_data = extract_coherence_phase_lag(pdf_paths)
            other_sections_data = extract_other_sections(pdf_paths)
            
            # Merge the data
            merged_data = merge_data(coherence_phase_lag_data + other_sections_data)
            
            # Create output directory with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = UPLOAD_DIR / timestamp
            output_dir.mkdir(exist_ok=True)
            
            # Save CSV files
            csv_output_path = output_dir / "extracted_data.csv"
            save_to_csv(merged_data, str(csv_output_path))
            
            # For each section CSV, apply calculations
            for section in merged_data.keys():
                section_csv_path = output_dir / f"{section}.csv"
                section_type = section
                apply_all_calculations(str(section_csv_path), section_type)
            
            # Create the combined document
            doc_output_path = output_dir / "eeg_analysis_summary.docx"
            create_combined_document(
                glob.glob(str(output_dir / "*.csv")),
                str(doc_output_path)
            )

            # AI interpretation DOCX (same metrics pattern as summary; requires GEMINI_API_KEY)
            csv_list = glob.glob(str(output_dir / "*.csv"))
            interpretation_path = output_dir / "eeg_ai_interpretation.docx"
            diagnostics_json_path = output_dir / "eeg_interpretation_diagnostics.json"
            interpretation_ok, interpretation_err = try_generate_ai_interpretation_docx(
                str(interpretation_path),
                csv_list,
                diagnostics_json_path=str(diagnostics_json_path),
            )
            
            # Zip the output directory
            zip_path = str(output_dir) + ".zip"
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, start=output_dir)
                        zipf.write(file_path, arcname)
            
            # Upload the zip file to DigitalOcean Spaces
            try:
                s3_url, unique_filename = upload_to_s3(zip_path)
            except ValueError as e:
                # Clean up local files before raising error
                shutil.rmtree(output_dir)
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                raise HTTPException(status_code=500, detail=str(e))
            
            # Upload the docx file to DigitalOcean Spaces
            try:
                doc_s3_url, doc_unique_filename = upload_to_s3(str(doc_output_path))
            except ValueError as e:
                # Clean up local files before raising error
                shutil.rmtree(output_dir)
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                raise HTTPException(status_code=500, detail=str(e))

            interpretation_url = None
            interpretation_filename = None
            if interpretation_ok and interpretation_path.exists():
                try:
                    interpretation_url, interpretation_filename = upload_to_s3(
                        str(interpretation_path)
                    )
                except ValueError as e:
                    shutil.rmtree(output_dir)
                    if os.path.exists(zip_path):
                        os.remove(zip_path)
                    raise HTTPException(status_code=500, detail=str(e))
            
            # Clean up local files
            shutil.rmtree(output_dir)
            os.remove(zip_path)

            # Return the DigitalOcean Spaces CDN URLs
            response_body = {
                "message": "File processed and uploaded successfully to DigitalOcean Spaces",
                "url": s3_url,
                "filename": unique_filename,
                "doc_url": doc_s3_url,
                "doc_filename": doc_unique_filename,
                "interpretation_url": interpretation_url,
                "interpretation_filename": interpretation_filename,
            }
            if interpretation_err:
                response_body["interpretation_note"] = interpretation_err
            return JSONResponse(status_code=200, content=response_body)
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/extraction/v2")
async def extract_data_v2(
    pdf1: UploadFile = File(...),
    pdf2: UploadFile = File(...)
):
    """
    Upload 2 PDFs, extract data, perform all calculations, and return HTML format only.
    This is the v2 endpoint that provides HTML output.
    """
    # Validate file types
    for pdf in [pdf1, pdf2]:
        if not pdf.filename.lower().endswith('.pdf'):
            raise HTTPException(
                status_code=400, 
                detail=f"File {pdf.filename} is not a PDF"
            )
    
    try:
        # Create a temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_paths = []
            
            # Save uploaded files
            for pdf in [pdf1, pdf2]:
                file_path = os.path.join(temp_dir, pdf.filename)
                with open(file_path, "wb") as buffer:
                    shutil.copyfileobj(pdf.file, buffer)
                pdf_paths.append(file_path)
            
            # Process the PDFs
            coherence_phase_lag_data = extract_coherence_phase_lag(pdf_paths)
            other_sections_data = extract_other_sections(pdf_paths)
            
            # Merge the data
            merged_data = merge_data(coherence_phase_lag_data + other_sections_data)
            
            # Create output directory with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = UPLOAD_DIR / timestamp
            output_dir.mkdir(exist_ok=True)
            
            # Save CSV files
            csv_output_path = output_dir / "extracted_data.csv"
            save_to_csv(merged_data, str(csv_output_path))
            
            # For each section CSV, apply calculations
            for section in merged_data.keys():
                section_csv_path = output_dir / f"{section}.csv"
                section_type = section
                apply_all_calculations(str(section_csv_path), section_type)
            
            # Create the combined HTML document
            html_output_path = output_dir / "eeg_analysis_summary.html"
            create_combined_html_document(
                glob.glob(str(output_dir / "*.csv")),
                str(html_output_path)
            )
            
            # Zip the output directory
            zip_path = str(output_dir) + ".zip"
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, start=output_dir)
                        zipf.write(file_path, arcname)
            
            # Upload the zip file to DigitalOcean Spaces
            try:
                s3_url, unique_filename = upload_to_s3(zip_path)
            except ValueError as e:
                # Clean up local files before raising error
                shutil.rmtree(output_dir)
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                raise HTTPException(status_code=500, detail=str(e))
            
            # Upload the HTML file to DigitalOcean Spaces
            try:
                html_s3_url, html_unique_filename = upload_to_s3(str(html_output_path))
            except ValueError as e:
                # Clean up local files before raising error
                shutil.rmtree(output_dir)
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                raise HTTPException(status_code=500, detail=str(e))
            
            # Clean up local files
            shutil.rmtree(output_dir)
            os.remove(zip_path)

            # Return the DigitalOcean Spaces CDN URLs
            return JSONResponse(
                status_code=200,
                content={
                    "message": "File processed and uploaded successfully to DigitalOcean Spaces (v2 - HTML format)",
                    "url": s3_url,
                    "filename": unique_filename,
                    "html_url": html_s3_url,
                    "html_filename": html_unique_filename
                }
            )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/extraction/cswl")
async def extract_data_cswl(
    cswl1: UploadFile = File(...),
    cswl2: UploadFile = File(...),
):
    """
    Upload 2 CSWL files and return a DOCX report in the same style as /extraction.
    """
    for f in [cswl1, cswl2]:
        if not f.filename or not f.filename.lower().endswith(".cswl"):
            raise HTTPException(status_code=400, detail=f"File {f.filename} is not a CSWL file")

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir_path = Path(temp_dir)
            cswl_paths: list[Path] = []

            for f in [cswl1, cswl2]:
                safe_name = secure_cswl_filename(f.filename)
                file_bytes = await f.read()
                if len(file_bytes) > MAX_CSWL_UPLOAD_BYTES:
                    raise HTTPException(status_code=413, detail=f"File {safe_name} exceeds max size limit")
                save_path = temp_dir_path / safe_name
                save_path.write_bytes(file_bytes)
                cswl_paths.append(save_path)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = UPLOAD_DIR / f"cswl_{timestamp}"
            output_dir.mkdir(exist_ok=True)

            doc_output_path = process_cswl_pair_to_output_dir(
                cswl_paths[0], cswl_paths[1], output_dir
            )

            zip_path = str(output_dir) + ".zip"
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, start=output_dir)
                        zipf.write(file_path, arcname)

            try:
                s3_url, unique_filename = upload_to_s3(zip_path)
                doc_s3_url, doc_unique_filename = upload_to_s3(str(doc_output_path))
            except ValueError as e:
                shutil.rmtree(output_dir)
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                raise HTTPException(status_code=500, detail=str(e))

            shutil.rmtree(output_dir)
            os.remove(zip_path)

            return JSONResponse(
                status_code=200,
                content={
                    "message": "CSWL files processed and uploaded successfully",
                    "url": s3_url,
                    "filename": unique_filename,
                    "doc_url": doc_s3_url,
                    "doc_filename": doc_unique_filename,
                },
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def root():
    """
    Root endpoint that returns API information
    """
    return {
        "message": "Welcome to PDF Data Extraction API",
        "endpoints": {
            "/ui": "GET - Web UI (upload PDFs, download result links)",
            "/extraction": "POST - Upload and process PDF files (returns DOCX)",
            "/extraction/v2": "POST - Upload and process PDF files (returns HTML format only)",
            "/extraction/cswl": "POST - Upload and process CSWL files (returns DOCX)",
            "/docs": "GET - API documentation"
        }
    }

@app.post("/calculation")
async def calculation(
    csv_file: UploadFile = File(...),
    section_type: str = File(...)
):
    """
    Accepts a CSV file and section type, returns band-wise absolute sums.
    """
    # Save uploaded CSV to a temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
        tmp.write(await csv_file.read())
        tmp_path = tmp.name

    try:
        summary = bandwise_absolute_sum(tmp_path, section_type)
        return {"summary": summary}
    finally:
        os.remove(tmp_path)

def create_combined_document(csv_files, output_path):
    """
    Create a single document combining all CSV files with metrics analysis
    """
    doc = Document()
    # Add custom heading at the top
    eeg_heading = doc.add_paragraph()
    run = eeg_heading.add_run("EEG//")
    run.bold = True
    run.font.size = Pt(18)
    # Default color (do not set run.font.color)
    doc.add_heading("Band-wise Metrics Analysis: ", level=1)

    # Define the desired order of sections
    section_order = [
        "Z Scored FFT Absolute Power",
        "Z Scored FFT Power Ratio",
        "Z Scored Peak Frequency",
        "Z Scored FFT Coherence",
        "Z Scored FFT Phase Lag"
    ]
    # Map section names to filenames
    section_to_file = {os.path.basename(f).replace('.csv', ''): f for f in csv_files}

    # Metrics to bullet
    bullet_metrics = [
        "Absolute Sum:",
        "Average Absolute Value:",
        "Delta:",
        "Percent Change:",
        "Total Rows:",
        '"Normalize = Yes":',
        '"Normalize = No":',
        '"Normalize = NS":',
    ]

    # Iterate in the desired order
    for section in section_order:
        file_path = section_to_file.get(section)
        if not file_path:
            continue
        section_name = section
        df = pd.read_csv(file_path)
        metrics_text = compute_subsection_bandwise_metrics(
            df,
            subsection_col="Subsection",
            band_col="Band", 
            set1_col="T1 Z",
            set2_col="T2 Z",
            normalize_col="Normalize"
        )
        section_heading = doc.add_heading(section_name, level=2)
        for block in metrics_text.strip().split("\n Subsection: "):
            if block.strip():
                if not block.startswith("Subsection:"):
                    block = "Subsection: " + block
                lines = block.strip().split("\n", 1)
                subsection_title = lines[0].replace("Subsection: ", "").strip()
                content = lines[1] if len(lines) > 1 else ""
                para = doc.add_paragraph()
                run = para.add_run(f"Subsection: {subsection_title}")
                run.bold = True
                # Add each line, prepending a bullet if it's a metric
                for line in content.strip().split("\n"):
                    line_strip = line.strip()
                    if any(line_strip.lstrip().startswith(metric) for metric in bullet_metrics):
                        p = doc.add_paragraph(f"• {line_strip}", style='Normal')
                        p.paragraph_format.left_indent = Inches(0.4)
                    elif line_strip:
                        doc.add_paragraph(line_strip, style='Normal')
        doc.add_paragraph()
    doc.save(output_path)
    print(f"Combined analysis saved to {output_path}")

def process_all_csv_files(input_directory, output_doc="combined_analysis.docx"):
    """
    Process all CSV files in the input directory and create a combined document
    """
    # Get all CSV files in the directory
    csv_files = glob.glob(os.path.join(input_directory, "*.csv"))
    
    if not csv_files:
        print(f"No CSV files found in {input_directory}")
        return
    
    # Create the combined document
    create_combined_document(csv_files, output_doc)

def main():
    # Example usage
    input_directory = "uploads"  # Directory containing your CSV files
    output_doc = "eeg_analysis_summary.docx"
    
    process_all_csv_files(input_directory, output_doc)

def bandwise_absolute_sum(csv_path: str, section_type: str) -> dict:
    """Calculate band-wise absolute sums for a given CSV file"""
    df = pd.read_csv(csv_path)
    summary = {
        'section_type': section_type,
        'subsections': {}
    }
    
    for subsection in df['Subsection'].unique():
        sub_df = df[df['Subsection'] == subsection]
        summary['subsections'][subsection] = {}
        
        for band in sub_df['Band'].unique():
            band_df = sub_df[sub_df['Band'] == band]
            t1_sum = band_df['T1 Z'].abs().sum()
            t2_sum = band_df['T2 Z'].abs().sum()
            summary['subsections'][subsection][band] = {
                'T1_sum': t1_sum,
                'T2_sum': t2_sum
            }
    
    return summary

def apply_all_calculations(csv_path: str, section_type: str) -> None:
    """Apply all calculations to a CSV file and save results"""
    df = pd.read_csv(csv_path)
    
    # Add section type as a column
    df['Section_Type'] = section_type
    
    # Calculate metrics
    for subsection in df['Subsection'].unique():
        sub_df = df[df['Subsection'] == subsection]
        for band in sub_df['Band'].unique():
            band_df = sub_df[sub_df['Band'] == band]
            
            # Calculate absolute sums
            t1_sum = band_df['T1 Z'].abs().sum()
            t2_sum = band_df['T2 Z'].abs().sum()
            
            # Calculate averages
            t1_avg = band_df['T1 Z'].abs().mean()
            t2_avg = band_df['T2 Z'].abs().mean()
            
            # Calculate delta and percent change
            delta = abs(t1_avg - t2_avg)
            percent_change = (delta / t1_avg) * 100 if t1_avg != 0 else 0
            
            # Update the DataFrame with calculations
            mask = (df['Subsection'] == subsection) & (df['Band'] == band)
            df.loc[mask, 'T1_Sum'] = t1_sum
            df.loc[mask, 'T2_Sum'] = t2_sum
            df.loc[mask, 'Delta'] = delta
            df.loc[mask, 'Percent_Change'] = percent_change
    
    # Save the updated DataFrame
    df.to_csv(csv_path, index=False)

def create_combined_html_document(csv_files, output_path):
    """
    Create a single HTML document combining all CSV files with metrics analysis.
    Follows the HTML format guide for Dart Editor.
    Returns only the body content without HTML structure.
    """
    html_content = []
    
    # Add custom heading at the top - EEG//
    html_content.append("<p><strong>EEG//</strong></p>")
    html_content.append("<p></p>")
    
    # Main document title
    html_content.append("<h1>Band-wise Metrics Analysis</h1>")
    html_content.append("<p></p>")
    
    # Define the desired order of sections
    section_order = [
        "Z Scored FFT Absolute Power",
        "Z Scored FFT Power Ratio",
        "Z Scored Peak Frequency",
        "Z Scored FFT Coherence",
        "Z Scored FFT Phase Lag"
    ]
    
    # Map section names to filenames
    section_to_file = {os.path.basename(f).replace('.csv', ''): f for f in csv_files}
    
    # Metrics to bullet
    bullet_metrics = [
        "Absolute Sum:",
        "Average Absolute Value:",
        "Delta:",
        "Percent Change:",
        "Total Rows:",
        '"Normalize = Yes":',
        '"Normalize = No":',
        '"Normalize = NS":',
    ]
    
    # Iterate in the desired order
    for section in section_order:
        file_path = section_to_file.get(section)
        if not file_path:
            continue
        
        section_name = section
        df = pd.read_csv(file_path)
        metrics_text = compute_subsection_bandwise_metrics(
            df,
            subsection_col="Subsection",
            band_col="Band", 
            set1_col="T1 Z",
            set2_col="T2 Z",
            normalize_col="Normalize"
        )
        
        # Section heading (H2)
        html_content.append(f"<h2>{html.escape(section_name)}</h2>")
        
        # Process each subsection block
        for block in metrics_text.strip().split("\n Subsection: "):
            if block.strip():
                if not block.startswith("Subsection:"):
                    block = "Subsection: " + block
                
                lines = block.strip().split("\n", 1)
                subsection_title = lines[0].replace("Subsection: ", "").strip()
                content = lines[1] if len(lines) > 1 else ""
                
                # Subsection header using <pre> tag (as per guide)
                html_content.append(f"<pre>Subsection: {html.escape(subsection_title)}</pre>")
                
                # Process content lines - parse band by band
                current_band = None
                in_set1 = False
                in_set2 = False
                in_differences = False
                in_normalize = False
                
                for line in content.strip().split("\n"):
                    line_strip = line.strip()
                    
                    if not line_strip or line_strip == "__________________________________________________":
                        if current_band:
                            html_content.append("<p></p>")  # Empty line after band
                        continue
                    
                    # Check if it's a Band line
                    if line_strip.startswith("Band:"):
                        current_band = line_strip.replace("Band:", "").strip()
                        html_content.append(f"<p><strong>Band: {html.escape(current_band)}</strong></p>")
                        in_set1 = False
                        in_set2 = False
                        in_differences = False
                        in_normalize = False
                    
                    # Check for section headers
                    elif line_strip.startswith("Set 1"):
                        in_set1 = True
                        in_set2 = False
                        html_content.append(f"<p><strong>{html.escape(line_strip)}</strong></p>")
                    
                    elif line_strip.startswith("Set 2"):
                        in_set1 = False
                        in_set2 = True
                        html_content.append(f"<p><strong>{html.escape(line_strip)}</strong></p>")
                    
                    elif line_strip.startswith("Differences:"):
                        in_differences = True
                        in_normalize = False
                        html_content.append(f"<p><strong>{html.escape(line_strip)}</strong></p>")
                    
                    elif line_strip.startswith("Normalize Counts:"):
                        in_normalize = True
                        in_differences = False
                        html_content.append(f"<p><strong>{html.escape(line_strip)}</strong></p>")
                    
                    # Check if it's a metric line (bullet format)
                    elif any(line_strip.lstrip().startswith(metric) for metric in bullet_metrics):
                        # Use bullet format for metrics as per guide
                        html_content.append(f"<p>• {html.escape(line_strip)}</p>")
                    
                    elif line_strip:
                        # Regular text line - use span with br for less spacing
                        html_content.append(f"<span>{html.escape(line_strip)}</span><br/>")
                
                html_content.append("<p></p>")  # Empty line between subsections
        
        html_content.append("<p></p>")  # Empty line between sections
    
    # Write HTML to file (only body content, no HTML structure)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(html_content))
    
    print(f"Combined HTML analysis saved to {output_path}")

@app.get("/download/{timestamp}/{filename}")
async def download_file(timestamp: str, filename: str):
    file_path = UPLOAD_DIR / timestamp / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(str(file_path), filename=filename)

if __name__ == "__main__":
    main() 